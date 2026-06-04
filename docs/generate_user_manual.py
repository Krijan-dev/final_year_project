"""One-off generator for docs/User_Manual_In_A_Nutshell.docx — run: python docs/generate_user_manual.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "User_Manual_In_A_Nutshell.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Life Pattern Tracker", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("User manual — in a nutshell")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].italic = True
    doc.add_paragraph()

    add_para(
        doc,
        "Life Pattern Tracker helps you understand phone use, build habits, track steps and sleep, "
        "and get personalised insights. This guide explains what each part of the app does and how to use it.",
    )

    add_heading(doc, "1. Getting started", 1)
    add_bullets(
        doc,
        [
            "Install the app from your tester link or Play Store build.",
            "Open the app and swipe through the welcome screens, then tap Get started.",
            "Sign in with your email and password, or create a new account.",
            "Forgot password? On the sign-in screen, tap Forgot password and follow the email code steps.",
            "After sign-in, you land on the main app with five tabs along the bottom.",
        ],
    )
    add_para(doc, "Account menu: tap your profile circle (top-right on most screens) to open Account.")

    add_heading(doc, "2. Bottom navigation — five main tabs", 1)
    add_para(doc, "Use the bar at the bottom to move between sections:")

    tabs = [
        (
            "Home",
            "Your daily dashboard. Shows screen time, focus score, wellness-style scores, "
            "quick calculated insights, and (when set up) steps from Health Connect. "
            "Pull down to refresh. If Usage access is off, a card here explains how to turn it on.",
        ),
        (
            "Time",
            "Screen time in detail: total time today, charts, app list, and per-app daily limits. "
            "Matches your phone’s Usage access data (not cloud). Pull down to refresh. "
            "Set limits on apps to get notified when you exceed them.",
        ),
        (
            "Habits",
            "Create habits, log them daily, track weekly progress, streaks, and mood. "
            "Tap a habit for history and details. Add habits with the + button. "
            "Pull down to refresh.",
        ),
        (
            "Insights",
            "Health risk score, wellness scores, smart recommendations (from your usage and habits), "
            "and optional AI insights when configured. Pull down to refresh after you have usage or habit data.",
        ),
        (
            "Health",
            "Steps today, sleep last night, wellness score, 7-day step trend, and tips. "
            "Data comes from Health Connect (Samsung Health, Google Fit, Fitbit, etc.). "
            "See section 5 for setup. Tap Refresh after syncing your fitness app.",
        ),
    ]
    for name, desc in tabs:
        add_heading(doc, name, 2)
        add_para(doc, desc)

    add_heading(doc, "3. Floating chat (green button)", 1)
    add_bullets(
        doc,
        [
            "On every main tab, a green chat button sits at the bottom-right.",
            "Tap it to open the assistant: ask about screen time, habits, health, or app help.",
            "When signed in, you can also request live human support from the chat screen.",
            "For urgent mental health help, use Account → Crisis support (Lifeline 13 11 14, Emergency 000).",
        ],
    )

    add_heading(doc, "4. Account — settings, permissions, backup", 1)
    add_para(doc, "Open Account from the profile icon. Main areas:")

    account_sections = [
        (
            "Profile",
            "Shows your signed-in email.",
        ),
        (
            "Account & security",
            "Reset password: sends a code to your email.",
        ),
        (
            "Permissions & data sources",
            "Usage access — required for screen time (Time tab and Home). "
            "Shows On/Off; tap to open Android settings. After granting, tap “I granted permission — check again”. "
            "Health Connect — opens the Health screen for steps and sleep. "
            "Screen time & app limits — shortcut to the Time tab.",
        ),
        (
            "Cloud backup (when signed in)",
            "Back up to cloud now — saves habits to your account (screen time stays on this phone only). "
            "Restore habits from cloud — downloads habits; does not replace local screen time.",
        ),
        (
            "Appearance",
            "System, Light, or Dark theme.",
        ),
        (
            "Help & support",
            "Live support chat — how to use the green chat button. Crisis support — helpline numbers.",
        ),
        (
            "Log out / Delete account",
            "Log out ends your session on this device. Delete account permanently removes your account and cloud data.",
        ),
    ]
    for title, body in account_sections:
        add_heading(doc, title, 2)
        add_para(doc, body)

    add_heading(doc, "5. Health Connect — steps and sleep", 1)
    add_bullets(
        doc,
        [
            "Install or update Health Connect from the Play Store (on Android 14+ it may be in Settings).",
            "In Samsung Health, Google Fit, Fitbit, Garmin, etc., turn on sharing to Health Connect.",
            "In Health Connect, allow Life Pattern Tracker to read Steps and Sleep.",
            "Open the Health tab and tap Grant access or Check again if needed.",
            "If a fitness app is detected, tap Open Samsung Health (or similar) to enable sharing, then Refresh.",
            "Under the status bar you may see Last updated … from [app] — that means Health Connect has recent data.",
            "If you see data may be stale, open your fitness app, wait for sync, then tap Refresh on Health.",
        ],
    )

    add_heading(doc, "6. Screen time (Usage access)", 1)
    add_bullets(
        doc,
        [
            "Screen time uses Android Usage access — the same type of data your phone’s Digital Wellbeing uses.",
            "The app does not block launch until you grant it; prompts appear on Home and Time when needed.",
            "Grant path varies by phone; the app shows device-specific hints (e.g. Samsung: Settings → Apps → Special access → Usage access).",
            "Enable Life Pattern Tracker in that list, return to the app, and tap check again.",
            "Usage access status (On/Off) is shown under Account → Permissions, not as a banner on Home.",
        ],
    )

    add_heading(doc, "7. Habits — quick how-to", 1)
    add_bullets(
        doc,
        [
            "Add a habit: Habits tab → + → name, schedule, and optional reminder.",
            "Log today: tap the habit or use quick log on the Habits screen.",
            "Mood: log how you feel from the mood section on Habits.",
            "Weekly progress card shows how you are doing across all habits.",
            "Habits sync to the cloud when you use Back up; screen time does not.",
        ],
    )

    add_heading(doc, "8. Insights — what the numbers mean", 1)
    add_bullets(
        doc,
        [
            "Health risk score — summary score from usage, habits, sleep, and steps when available.",
            "Wellness scores — breakdown areas (e.g. balance, focus) calculated on your device.",
            "Smart recommendations — rule-based tips from your patterns (no AI required).",
            "AI insights — extra text when the server AI is configured; may take a moment to load.",
        ],
    )

    add_heading(doc, "9. Tips & troubleshooting", 1)
    table_data = [
        ("Screen time is zero or wrong", "Turn on Usage access in Account; open Time tab and pull to refresh. Compare with your phone’s built-in screen time after granting access."),
        ("Steps or sleep missing", "Set up Health Connect (section 5). Open your fitness app once, then Refresh on Health."),
        ("Steps lower than Samsung Health", "Fitness app must sync to Health Connect first; our app reads Health Connect only."),
        ("Habits lost on new phone", "Sign in → Account → Restore habits from cloud. Re-enable Usage access for screen time on the new device."),
        ("Chat does not connect to a person", "You must be signed in; use the green chat button and choose live support."),
        ("App looks outdated", "Install the latest tester APK from your Firebase / invite link."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Problem"
    hdr[1].text = "What to try"
    for prob, fix in table_data:
        row = table.add_row().cells
        row[0].text = prob
        row[1].text = fix

    doc.add_paragraph()
    add_para(doc, "Version: 1.0 · Platform: Android · Life Pattern Tracker", bold=False)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
