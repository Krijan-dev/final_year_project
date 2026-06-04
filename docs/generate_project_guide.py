"""Build docs/PROJECT_COMPLETE_GUIDE.docx — run: python docs/generate_project_guide.py"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parent / "PROJECT_COMPLETE_GUIDE.docx"


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    run.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Life Pattern Tracker", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Complete project guide — app, Render API, website, docs")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    p(
        doc,
        "This guide explains how the Android app, Node API on Render, MongoDB Atlas, "
        "and the Vercel admin website work together. For day-to-day app use, see User_Manual_In_A_Nutshell.docx.",
    )

    h(doc, "1. System overview")
    p(
        doc,
        "Flutter app → HTTPS → Render (server/) → MongoDB Atlas. "
        "The Next.js website (separate repo, Vercel) uses the same Render URL. "
        "The app never talks to MongoDB directly.",
    )
    bullets(
        doc,
        [
            "Screen time: read on the phone via Usage Access (not downloaded from cloud for display).",
            "Habits/mood/logs: stored in Hive locally; backed up to habit_snapshots via API.",
            "Health: read from Health Connect on the device (Samsung Health, Fitbit, etc. via HC sync).",
            "AI: Gemini key in flutter.env (app → Google, not via Render).",
        ],
    )

    h(doc, "2. Repository layout", 1)
    bullets(
        doc,
        [
            "lib/ — Flutter UI and business logic",
            "android/ — Usage stats, Health Connect bridge, OEM settings helpers",
            "server/ — Express API (deploy root on Render)",
            "docs/ — all documentation (see section 9)",
            "scripts/ — sync-env, Firebase release, flutter.env sync",
        ],
    )

    h(doc, "3. Flutter app", 1)
    h(doc, "Bottom tabs", 2)
    for tab, desc in [
        ("Home", "Dashboard: screen time, wellness scores, quick insights."),
        ("Time", "Detailed screen time, charts, per-app limits."),
        ("Habits", "Weekly habits, mood, logging; Hive + cloud backup."),
        ("Insights", "Risk/wellness scores, recommendations, optional AI."),
        ("Health", "Steps, sleep, trends from Health Connect."),
    ]:
        p(doc, f"{tab}: {desc}")

    h(doc, "Key providers", 2)
    bullets(
        doc,
        [
            "authProvider — remote login/register when API_BASE_URL is set",
            "usageProvider — Usage Access permission and today’s stats",
            "habitTrackerProvider — local habits; sync via CloudSyncService",
            "dashboardProvider / insightsProvider — aggregated metrics",
        ],
    )

    h(doc, "Cloud sync rules", 2)
    bullets(
        doc,
        [
            "Sign-in: pull habits only if local empty; push current week; screen time stays device-only.",
            "Account → Back up: uploads usage history + habits to MongoDB.",
            "Account → Restore habits: downloads habits; does not replace phone screen time.",
        ],
    )

    h(doc, "4. Render API (server/)", 1)
    p(doc, "Deploy: Render web service, root directory server, npm install / npm start.")
    p(doc, "Required env: MONGODB_URI. Optional: ADMIN_EMAIL, ADMIN_PASSWORD, SMTP_* for emails.")
    p(doc, "Health: GET /health → ok and mongo true.")
    h(doc, "Main API groups", 2)
    bullets(
        doc,
        [
            "Auth: send-verification, verify-email, register, login, logout, forgot/reset password",
            "User data (Bearer): usage-days, habit-snapshot (own email only)",
            "Support: user conversations/messages; admin support routes",
            "Admin: login, list users, usage-days, habit snapshot, stats, crisis flags",
        ],
    )

    h(doc, "5. MongoDB Atlas collections", 1)
    bullets(
        doc,
        [
            "users — email, passwordHash, sessionToken",
            "usagedays — per-day screen time JSON",
            "habitsnapshots — weekKey + habits + mood + logs",
            "supportconversations / supportmessages — live chat",
            "crisisflags — safety alerts from chat",
        ],
    )

    h(doc, "6. Connect the app to Render", 1)
    bullets(
        doc,
        [
            "Set API_BASE_URL=https://your-service.onrender.com in root .env",
            "Run scripts/sync-env.ps1 if syncing to server/website clones",
            "Rebuild app: run_dev.ps1 or flutter run --dart-define-from-file=flutter.env",
            "Test /health in a browser before testing login on the phone",
        ],
    )

    h(doc, "7. Website (Vercel)", 1)
    p(
        doc,
        "Repository: LifePatternAI_Website (GitHub). Hosted on Vercel. "
        "Uses NEXT_PUBLIC_API_BASE_URL pointing at the same Render service. "
        "Admin pages call /api/v1/admin/* routes. Does not connect to MongoDB directly.",
    )
    bullets(
        doc,
        [
            "Landing and marketing pages",
            "Admin login (ADMIN_EMAIL / ADMIN_PASSWORD on API)",
            "User list and per-user usage/habit views",
            "Support chat and Safety alerts pages",
        ],
    )

    h(doc, "8. Environment variables", 1)
    bullets(
        doc,
        [
            "Root .env — API_BASE_URL, GEMINI_API_KEY, MONGODB_URI, ADMIN_*, SMTP_*",
            "scripts/sync-env.ps1 — copies to server/.env and website .env.local",
            "flutter.env — GEMINI + API only for release APKs",
            "Render dashboard — MONGODB_URI and production secrets",
            "Vercel dashboard — NEXT_PUBLIC_API_BASE_URL",
        ],
    )

    h(doc, "9. Docs folder map", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "File"
    table.rows[0].cells[1].text = "Purpose"
    rows = [
        ("PROJECT_COMPLETE_GUIDE", "This guide (MD + DOCX)"),
        ("User_Manual_In_A_Nutshell.docx", "End-user app manual"),
        ("HOW_TO_RUN.md", "Local dev setup"),
        ("DEPLOY_API_CLOUD.md", "Render deployment"),
        ("VERCEL_WEBSITE_GUIDE.md", "Website on Vercel"),
        ("ENV_SETUP.md", "Env sync script"),
        ("MONGODB.md", "API + collections"),
        ("SUPPORT_CHAT.md", "Live support flow"),
        ("firebase_testing_quickstart.md", "Tester APK"),
        ("EMAIL_VERIFICATION.md / PASSWORD_RESET.md", "Auth email flows"),
    ]
    for a, b in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b

    h(doc, "10. Troubleshooting", 1)
    bullets(
        doc,
        [
            "Login fails — check API_BASE_URL and Render /health; rebuild app after env change",
            "Screen time zero — Usage access in Account, not cloud",
            "Steps missing — Health Connect + fitness app sharing; latest APK",
            "Website empty — Vercel env URL; admin login; cold Render start",
        ],
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
