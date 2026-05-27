"""Generate docs/DATABASE_AND_WEBSITE_PLAN.docx for GitHub / project documentation."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "docs" / "DATABASE_AND_WEBSITE_PLAN.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Life Pattern Tracker", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Database, Cloud Sync & Admin Website — Implementation Plan")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Final year project · Flutter Android app + Node.js API + MongoDB + future admin dashboard"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading(doc, "1. Executive summary", 1)
    doc.add_paragraph(
        "You already have the foundation for a cloud database: MongoDB Atlas (free) stores user "
        "accounts (hashed passwords) and daily screen-usage JSON via the Node.js API in the server/ "
        "folder. The Flutter app can sync usage when API_BASE_URL is configured."
    )
    doc.add_paragraph(
        "To support a website that displays all user data, you should: (1) host MongoDB on Atlas, "
        "(2) host the API on a free/low-cost platform, (3) extend the API to sync habits, mood, and "
        "activity logs (currently only on the phone in Hive), (4) build a separate admin web app "
        "that reads from the API with admin authentication — never expose raw passwords."
    )

    add_heading(doc, "2. What you already have in this project", 1)
    add_heading(doc, "2.1 MongoDB + API (server/)", 2)
    add_bullets(
        doc,
        [
            "MongoDB collections: users (email, passwordHash, sessionToken), usagedays (per user per day).",
            "Auth endpoints: POST /api/v1/auth/register, login, logout.",
            "Data endpoints: PUT/GET /api/v1/users/<email>/usage-days (Bearer token required).",
            "Passwords stored as salt + SHA-256 hash — never plain text.",
            "Documentation: docs/MONGODB.md, docs/HOW_TO_RUN.md.",
        ],
    )
    add_heading(doc, "2.2 Flutter app — what is stored where", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Data"
    hdr[1].text = "On device (Hive)"
    hdr[2].text = "In MongoDB (today)"
    rows = [
        ("User login / password", "Yes (local_auth_box)", "Only if you wire app to API*"),
        ("Daily screen time & apps", "Yes", "Yes, when API_BASE_URL set"),
        ("Weekly habits grid", "Yes", "No — needs new API"),
        ("Mood per day", "Yes", "No — needs new API"),
        ("Activity logs (Today’s Log)", "Yes", "No — needs new API"),
    ]
    for r in rows:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = r
    doc.add_paragraph(
        "*Note: Auth in the app currently uses local Hive (auth_storage_service.dart). "
        "auth_remote_service.dart exists for MongoDB but should be connected in auth_provider "
        "when API_BASE_URL is set so passwords and sessions live in the cloud."
    )

    add_heading(doc, "3. Recommended architecture", 1)
    doc.add_paragraph(
        "Keep the same pattern: Flutter → REST API → MongoDB. The website never connects to "
        "MongoDB directly (safer and easier to secure)."
    )
    doc.add_paragraph("Flow:")
    add_numbered(
        doc,
        [
            "User registers/logs in on the Android app → API creates user in MongoDB.",
            "App syncs usage, habits, mood, logs to API (authenticated with Bearer token).",
            "Admin website logs in as administrator → API returns aggregated/anonymised or per-user stats.",
            "MongoDB Atlas holds all persistent data.",
        ],
    )

    add_heading(doc, "4. Database schema to add (MongoDB)", 1)
    add_heading(doc, "4.1 Collection: habit_snapshots", 2)
    doc.add_paragraph(
        "One document per user per week (weekKey e.g. 2026-W21): habits array, moodDays array, "
        "updatedAt. Sync when user opens Habit tab or on background timer."
    )
    add_heading(doc, "4.2 Collection: activity_logs", 2)
    doc.add_paragraph(
        "Documents: userId, dateKey, activityKey, details, durationMinutes, timestamp. "
        "Or embed logs inside habit_snapshots for simpler queries."
    )
    add_heading(doc, "4.3 Collection: profiles (optional)", 2)
    doc.add_paragraph(
        "displayName, createdAt, lastSyncAt — no sensitive fields. Never store Gemini API keys in DB."
    )
    add_heading(doc, "4.4 Admin users (required for website)", 2)
    doc.add_paragraph(
        "Separate admin_accounts collection OR role field on users: role = user | admin. "
        "Admin login uses different JWT/session and can call GET /api/v1/admin/users, "
        "GET /api/v1/admin/stats — never return passwordHash to the browser."
    )

    add_heading(doc, "5. API work (server/) — phased plan", 1)
    add_heading(doc, "Phase 1 — Unify authentication (1–2 days)", 2)
    add_bullets(
        doc,
        [
            "Update Flutter auth_provider to call AuthRemoteService when API_BASE_URL is set.",
            "Fall back to local Hive only when API is offline (optional) or require online for register.",
            "Store session token in AuthTokenStore after login/register.",
        ],
    )
    add_heading(doc, "Phase 2 — Sync habits & mood (2–3 days)", 2)
    add_bullets(
        doc,
        [
            "PUT /api/v1/users/<email>/habit-week/<weekKey> — body: habits + moodDays JSON.",
            "GET same endpoint for restore on new device.",
            "Call from habit_tracker_provider after save.",
        ],
    )
    add_heading(doc, "Phase 3 — Sync activity logs (1–2 days)", 2)
    add_bullets(
        doc,
        [
            "POST batch endpoint for today’s log entries or full week sync.",
            "De-duplicate by userId + dateKey + activityKey + timestamp.",
        ],
    )
    add_heading(doc, "Phase 4 — Admin API for website (2–3 days)", 2)
    add_bullets(
        doc,
        [
            "POST /api/v1/admin/login (admin email + password, env-configured or DB).",
            "GET /api/v1/admin/users — list emails, lastSync, habit %, avg screen time (no passwords).",
            "GET /api/v1/admin/users/<email>/dashboard — usage + habits for charts.",
            "Enable CORS only for your website domain in production.",
        ],
    )

    add_heading(doc, "6. Admin website (what to build)", 1)
    doc.add_paragraph(
        "Purpose: view all registered users and their wellness data for demos, supervision, or research."
    )
    add_bullets(
        doc,
        [
            "Tech suggestions: Next.js or React + Vite (free hosting on Vercel/Netlify).",
            "Pages: Login, Users table, User detail (screen time chart, habit %, mood trend).",
            "Charts: Chart.js or Recharts — same metrics as app Dashboard/Insights.",
            "Auth: admin JWT in httpOnly cookie or localStorage (prefer cookie + HTTPS).",
        ],
    )
    doc.add_paragraph(
        "Important: This is an admin dashboard, not a public social site. Comply with privacy "
        "law (consent, data minimisation, ability to delete account)."
    )

    add_heading(doc, "7. Free & low-cost hosting options", 1)

    add_heading(doc, "7.1 Database hosting (recommended)", 2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = "Service", "Free tier", "Best for", "Notes"
    db_rows = [
        (
            "MongoDB Atlas",
            "M0 cluster — 512 MB",
            "This project (already documented)",
            "Best fit; drivers + Mongoose already used. Create cluster at mongodb.com/atlas.",
        ),
        (
            "Supabase (PostgreSQL)",
            "500 MB DB",
            "If you prefer SQL",
            "Would require rewriting server from Mongoose to SQL — not needed unless you want SQL.",
        ),
        (
            "Firebase Firestore",
            "Spark plan",
            "Mobile-first",
            "Different SDK; more rework than extending current API.",
        ),
    ]
    for r in db_rows:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = r

    add_heading(doc, "7.2 API (Node.js server/) hosting", 2)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text, h2[1].text, h2[2].text, h2[3].text = "Service", "Free tier", "Best for", "Notes"
    api_rows = [
        (
            "Render",
            "Free web service (sleeps after idle)",
            "Student projects",
            "Easy Node deploy from GitHub; set MONGODB_URI env var.",
        ),
        (
            "Railway",
            "Limited monthly credit",
            "Always-on small API",
            "Simple deploy; watch credit usage.",
        ),
        (
            "Fly.io",
            "Small free allowance",
            "Global edge",
            "Dockerfile deploy; good if you use docker-compose.yml.",
        ),
        (
            "Oracle Cloud Always Free",
            "VM always free",
            "Full control",
            "Run node + pm2 yourself; more setup, no sleep.",
        ),
    ]
    for r in api_rows:
        c = t2.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = r

    add_heading(doc, "7.3 Admin website hosting", 2)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    h3 = t3.rows[0].cells
    h3[0].text, h3[1].text, h3[2].text, h3[3].text = "Service", "Free tier", "Best for", "Notes"
    web_rows = [
        ("Vercel", "Hobby", "Next.js / React", "Connect GitHub repo; auto HTTPS."),
        ("Netlify", "Starter", "Static React/Vite", "Same as Vercel for static sites."),
        ("Cloudflare Pages", "Free", "Static frontend", "Fast CDN; API still on Render/Railway."),
        ("GitHub Pages", "Free", "Static only", "Cannot run Node API — frontend only."),
    ]
    for r in web_rows:
        c = t3.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = r

    add_heading(doc, "7.4 Recommended stack for your thesis/demo", 2)
    doc.add_paragraph(
        "MongoDB Atlas (database) + Render (API) + Vercel (admin website). "
        "Total cost: $0 for development and moderate demo traffic."
    )

    add_heading(doc, "8. Security checklist (mandatory)", 1)
    add_bullets(
        doc,
        [
            "Never store plain-text passwords — keep current hash approach or upgrade to bcrypt.",
            "Never commit .env, server/.env, or MongoDB URI to GitHub.",
            "Use HTTPS everywhere in production (Render/Vercel provide TLS).",
            "Admin routes protected; rate-limit login endpoints.",
            "Website must not display password hashes or session tokens.",
            "Atlas IP allowlist: restrict in production; 0.0.0.0/0 only for dev.",
            "Let users delete account + all data (GDPR-style).",
        ],
    )

    add_heading(doc, "9. Environment variables", 1)
    doc.add_paragraph("server/.env (on API host):")
    doc.add_paragraph("MONGODB_URI=<Atlas connection string>\nPORT=3000\nADMIN_EMAIL=you@university.edu\nADMIN_PASSWORD_HASH=<hash>")
    doc.add_paragraph("Flutter .env:")
    doc.add_paragraph("API_BASE_URL=https://your-api.onrender.com\nGEMINI_API_KEY=...")
    doc.add_paragraph("Admin website .env:")
    doc.add_paragraph("VITE_API_BASE_URL=https://your-api.onrender.com")

    add_heading(doc, "10. Timeline estimate", 1)
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text, t4.rows[0].cells[1].text, t4.rows[0].cells[2].text = (
        "Phase",
        "Effort",
        "Outcome",
    )
    for phase, effort, outcome in [
        ("Atlas + local API", "Done / 1 hour", "DB + auth + usage sync working"),
        ("Wire app auth to API", "1–2 days", "Cloud accounts"),
        ("Habits + mood sync", "2–3 days", "Full user data in MongoDB"),
        ("Admin API", "2–3 days", "Website can read all users"),
        ("Admin website UI", "3–5 days", "Charts and tables for thesis demo"),
        ("Deploy to Render + Vercel", "1 day", "Public URL for supervisor"),
    ]:
        c = t4.add_row().cells
        c[0].text, c[1].text, c[2].text = phase, effort, outcome

    add_heading(doc, "11. Next steps (action list)", 1)
    add_numbered(
        doc,
        [
            "Create MongoDB Atlas M0 cluster; add URI to server/.env.",
            "Run npm start in server/; verify GET /health.",
            "Set API_BASE_URL in Flutter .env; connect auth_provider to AuthRemoteService.",
            "Implement habit-week sync endpoints and Flutter sync calls.",
            "Build minimal admin website (user list + one user detail page).",
            "Deploy API to Render; deploy site to Vercel; update CORS and env URLs.",
        ],
    )

    add_heading(doc, "12. References in this repo", 1)
    add_bullets(
        doc,
        [
            "docs/MONGODB.md — Atlas setup and current schema",
            "docs/HOW_TO_RUN.md — full run guide",
            "server/src/index.js — API to extend",
            "lib/services/auth_remote_service.dart — Flutter API auth",
            "lib/services/usage_remote_service.dart — usage sync",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Document generated for Life Pattern Tracker — final year project.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
